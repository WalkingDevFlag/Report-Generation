import os
import csv
import docx
from docx2pdf import convert
from datetime import datetime
from pathlib import Path

def replace_text(element, replacements):
    """Replace text in document elements while preserving formatting"""
    if hasattr(element, 'text'):
        for key, value in replacements.items():
            if key in element.text:
                element.text = element.text.replace(key, value)
    
    if hasattr(element, 'runs'):
        for run in element.runs:
            for key, value in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)

def process_document(doc, replacements):
    """Process all elements in a Word document"""
    # Process paragraphs
    for paragraph in doc.paragraphs:
        replace_text(paragraph, replacements)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text(paragraph, replacements)
    
    # Process document headers and footers
    for section in doc.sections:
        for header in section.header.paragraphs:
            replace_text(header, replacements)
        for footer in section.footer.paragraphs:
            replace_text(footer, replacements)

def process_reports(csv_path, template_path, output_dir):
    """Generate reports from CSV data using Word template"""
    try:
        # Validate and resolve paths
        csv_path = Path(csv_path)
        template_path = Path(template_path)
        output_dir = Path(output_dir)
        
        if not csv_path.exists():
            raise FileNotFoundError(f"CSV file not found: {csv_path}")
        if not template_path.exists():
            raise FileNotFoundError(f"Template file not found: {template_path}")
        
        output_dir.mkdir(parents=True, exist_ok=True)

        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            required_columns = {'VIN Number', 'Street Address', 'City', 'ST ZIP Code'}
            if not required_columns.issubset(reader.fieldnames):
                missing = required_columns - set(reader.fieldnames)
                raise ValueError(f"CSV missing required columns: {', '.join(missing)}")

            for idx, row in enumerate(reader, 1):
                row = {k.strip(): v.strip() for k, v in row.items()}
                
                # Create unique filename
                vin = row.get('VIN Number', f"UNKNOWN_{idx}")
                safe_vin = "".join(c for c in vin if c.isalnum() or c in ('-', '_'))
                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                base_name = f"Report_{safe_vin}_{timestamp}"
                
                replacements = {
                    '[NUMBER]': f"{vin}-{idx:03d}",
                    '[DATE]': datetime.now().strftime("%d %b %Y"),
                    **{f'[{col}]': row.get(col, '') for col in reader.fieldnames}
                }

                # Process document
                doc = docx.Document(template_path)
                process_document(doc, replacements)

                # Save files
                docx_path = output_dir / f"{base_name}.docx"
                pdf_path = output_dir / f"{base_name}.pdf"
                
                doc.save(docx_path)
                convert(docx_path, pdf_path)
                
                if pdf_path.exists():
                    docx_path.unlink()  # Remove DOCX only if PDF exists
                else:
                    print(f"Warning: Failed to convert {docx_path} to PDF")

        print(f"✓ Successfully generated reports in: {output_dir.resolve()}")
        return True

    except Exception as e:
        print(f"✗ Critical error: {str(e)}")
        return False    

if __name__ == '__main__':
    # Hardcoded paths - change these as needed
    CSV_PATH = r"E:/Random Python Scripts/Report-Generation/Better Template/defects2.csv"
    TEMPLATE_PATH = r"E:/Random Python Scripts/Report-Generation/Better Template/template2.docx"
    OUTPUT_DIR = r"E:/Random Python Scripts/Report-Generation/Better Template/reports"
    
    process_reports(CSV_PATH, TEMPLATE_PATH, OUTPUT_DIR)