from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd


def set_default_font(doc, font_name="Calibri (Body)", font_size=12):
    """
    Set the default font and size for the document.
    """
    for style in doc.styles:
        if style.type == 1:  # Paragraph style
            font = style.font
            font.name = font_name
            font.size = Pt(font_size)


def add_list_bullet_style(doc):
    """
    Add a custom 'List Bullet' style to the document if it doesn't exist.
    """
    styles = doc.styles
    if "List Bullet" not in [style.name for style in styles]:
        style = styles.add_style("List Bullet", 1)  # 1 = Paragraph Style
        font = style.font
        font.name = "Calibri (Body)"
        font.size = Pt(12)


def clear_placeholder(doc, placeholder_text):
    """
    Clear paragraphs after the placeholder text to remove unwanted template content.
    """
    placeholder_index = None
    for i, paragraph in enumerate(doc.paragraphs):
        if placeholder_text in paragraph.text:
            placeholder_index = i
            break

    if placeholder_index is None:
        raise ValueError(f"Placeholder '{placeholder_text}' not found in the document")

    # Remove all paragraphs after the placeholder
    for _ in range(len(doc.paragraphs) - placeholder_index - 1):
        doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)

    return placeholder_index


def make_bold_and_resize(paragraph, font_size=14):
    """
    Make the given paragraph bold and resize the text.
    """
    run = paragraph.runs[0]  # Access the first run of the paragraph
    run.bold = True
    run.font.size = Pt(font_size)


def generate_dynamic_damage_report(template_path, output_path, data):
    """
    Generate a dynamic damage report based on the given template and data.
    """
    doc = Document(template_path)
    set_default_font(doc)  # Apply font settings to the entire document
    add_list_bullet_style(doc)  # Ensure 'List Bullet' style exists

    # Locate and clear the placeholder section
    placeholder_text = "Damage Breakdown:"
    placeholder_index = clear_placeholder(doc, placeholder_text)

    # Update the placeholder text and apply formatting
    placeholder_paragraph = doc.paragraphs[placeholder_index]
    placeholder_paragraph.text = "Damage Breakdown:"
    make_bold_and_resize(placeholder_paragraph, font_size=14)

    # Generate dynamic content based on the data
    for class_id, group in data.groupby('Class ID'):
        damage_type = group.iloc[0]['Damage Type']
        number_of_instances = len(group)

        # Add Class ID and Damage Type
        class_id_para = doc.add_paragraph(f"• Class ID {class_id}: {damage_type}")
        class_id_para.style.font.size = Pt(12)

        # Add Number of Instances
        doc.add_paragraph(f"  Number of Instances: {number_of_instances}").style = "List Bullet"

        # Add Mask Areas
        doc.add_paragraph("  Mask Areas:").style = "List Bullet"

        # Add each instance under Mask Areas, reset numbering for each damage type
        for idx, row in enumerate(group.iterrows(), start=1):
            bullet_text = f"    {damage_type} {idx}: {row[1]['Area']} sq. mm"
            instance_para = doc.add_paragraph(bullet_text, style="List Bullet")
            instance_para.paragraph_format.space_before = Pt(0)  # Shift + Enter effect

        # Add 2-line spacing between sections
        doc.add_paragraph("\n")

    # Save the document
    doc.save(output_path)


def prepare_report_data(csv_path, template_path, output_path):
    """
    Prepare data from the CSV file and generate a damage report.
    """
    # Read the CSV data
    df = pd.read_csv(csv_path)

    # Ensure the required columns are present
    required_columns = {'Class ID', 'Damage Type', 'Area'}
    if not required_columns.issubset(df.columns):
        raise ValueError(f"The CSV file must contain the following columns: {required_columns}")

    # Call the function to generate the report
    generate_dynamic_damage_report(template_path, output_path, df)


if __name__ == '__main__':
    csv_path = r'E:/Random Python Scripts/Report Generation/src/data.csv'
    template_path = r'E:/Random Python Scripts/Report Generation/src/template3.docx'
    output_path = r'E:/Random Python Scripts/Report Generation/src/Report.docx'

    prepare_report_data(csv_path, template_path, output_path)
