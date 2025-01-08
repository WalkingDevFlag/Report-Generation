from docx import Document
import pandas as pd
from docx.shared import Pt


def add_list_bullet_style(doc):
    """
    Add a custom 'List Bullet' style to the document if it doesn't exist.
    """
    styles = doc.styles
    if "List Bullet" not in [style.name for style in styles]:
        style = styles.add_style("List Bullet", 1)  # 1 = Paragraph Style
        font = style.font
        font.name = "Arial"
        font.size = Pt(12)


def generate_dynamic_damage_report(template_path, output_path, data):
    """
    Generate a dynamic damage report based on the given template and data.
    """
    doc = Document(template_path)
    add_list_bullet_style(doc)  # Ensure 'List Bullet' style exists

    # Locate the placeholder section in the template
    placeholder_index = None
    for i, paragraph in enumerate(doc.paragraphs):
        if "Damage Breakdown:" in paragraph.text:
            placeholder_index = i
            break

    if placeholder_index is None:
        raise ValueError("Placeholder 'Damage Breakdown:' not found in the template")

    # Clear existing placeholders in the document
    while len(doc.paragraphs) > placeholder_index + 1:
        doc.paragraphs.pop()

    # Generate dynamic content based on the data
    for class_id, group in data.groupby('Class ID'):
        damage_type = group.iloc[0]['Damage Type']
        number_of_instances = len(group)

        # Add Class ID and Damage Type
        doc.add_paragraph(f"• Class ID {class_id}: {damage_type}")
        doc.add_paragraph(f"• Number of Instances: {number_of_instances}")
        doc.add_paragraph("• Mask Areas:")

        # Add each instance under Mask Areas
        for idx, row in group.iterrows():
            doc.add_paragraph(f"  • {damage_type} {idx + 1}: {row['Area']} sq. mm", style="List Bullet")

        doc.add_paragraph()  # Add spacing between sections

    # Save the document
    doc.save(output_path)


def prepare_report_data(csv_path, template_path, output_path):
    """
    Prepare data from the CSV file and generate a damage report.
    """
    # Read the CSV data
    df = pd.read_csv(csv_path)

    # Ensure the required columns are present
    required_columns = {'Class ID', 'Damage Type', 'Area'}
    if not required_columns.issubset(df.columns):
        raise ValueError(f"The CSV file must contain the following columns: {required_columns}")

    # Call the function to generate the report
    generate_dynamic_damage_report(template_path, output_path, df)


if __name__ == '__main__':
    csv_path = r'E:/Random Python Scripts/Report Generation/src/data.csv'
    template_path = r'E:/Random Python Scripts/Report Generation/src/template3.docx'
    output_path = r'E:/Random Python Scripts/Report Generation/src/Report.docx'

    prepare_report_data(csv_path, template_path, output_path)
